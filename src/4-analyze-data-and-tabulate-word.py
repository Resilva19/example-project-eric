# Setup ------------------------------------------------------------------------

# Note: Some formatting is only done in LaTeX. You may wish to review the
# latex version of this script for additional / advanced formatting options.

# Load Libraries [i.e., packages]
# pip install python-dotenv pandas pyreadstat pyfixest scipy python-docx tabulate
#
# Unlike R's pacman::p_load, Python does not auto-install missing packages.
# Run the line above in your terminal once to install all dependencies.
#
# Key package equivalents:
#   R fixest::feols        -> Python pyfixest (pip install pyfixest)
#   R modelsummary         -> manual table building with pandas
#   R flextable + officer  -> python-docx (pip install python-docx)
#   R formattable          -> f-string formatting
#   R sjlabelled           -> pandas column rename / column_map dict

import os
import numpy as np
import pandas as pd
import pyreadstat
from dotenv import load_dotenv
from scipy import stats
import pyfixest as pf
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Load environment variables from .env file (see script 1 for detailed comments)
load_dotenv(".env")
data_dir = os.getenv("DATA_DIR")

# If the above is too complicated and you don't have coauthors you can just set
# data_dir manually by commenting out the two lines above and using:
# data_dir = "D:/Dropbox/example-project"


# Read in the data from the previous step --------------------------------------

# Read in the winsorized data.
# I found there are not many firms in the 60s so I am just going to start at 1970.
# pyreadstat.read_dta() returns a tuple: (DataFrame, metadata object).
regdata, meta = pyreadstat.read_dta(f"{data_dir}/regdata-R.dta")
regdata = regdata[["gvkey", "datadate", "calyear", "roa", "roa_lead_1",
                   "loss", "at", "mve", "rd", "FF12", "ff12num"]]

# Variable label map: used to rename columns in display tables.
# This is the Python equivalent of sjlabelled::var_labels().
var_labels = {
    "roa_lead_1": "ROA_{t+1}",
    "roa":        "ROA_t",
    "loss":       "LOSS",
    "rd":         "R&D",
    "at":         "TA",
    "mve":        "SIZE",
}


# Observations by Decade -------------------------------------------------------

## NOTE: see the latex version of this script for additional comments on
## each step

# Assign each observation to a decade bin.
# This is the Python equivalent of R's case_when() inside mutate().
def assign_decade(year):
    if 1970 <= year <= 1979:
        return "1970 - 1979"
    elif 1980 <= year <= 1989:
        return "1980 - 1989"
    elif 1990 <= year <= 1999:
        return "1990 - 1999"
    elif 2000 <= year <= 2009:
        return "2000 - 2009"
    elif 2010 <= year <= 2019:
        return "2010 - 2019"
    elif year >= 2020:
        return "2020 - 2022"

regdata["Year"] = regdata["calyear"].apply(assign_decade)

# Group by decade and compute frequency counts.
# Equivalent to R's group_by(Year) |> summarize(...)
t1 = (
    regdata
    .groupby("Year")
    .apply(lambda g: pd.Series({
        "Total Firms": f"{len(g):,.0f}",
        "Loss Firms":  f"{int(g['loss'].sum()):,.0f}",
        "Pct. Losses": f"{g['loss'].mean():.2%}",
    }))
    .reset_index()
)

# Add a total row (equivalent to R's bind_rows with a manually created totalrow)
totalrow = pd.DataFrame([{
    "Year":        "Total",
    "Total Firms": f"{len(regdata):,.0f}",
    "Loss Firms":  f"{int(regdata['loss'].sum()):,.0f}",
    "Pct. Losses": f"{regdata['loss'].mean():.2%}",
}])

# Bind the decade rows and the total row together
t1 = pd.concat([t1, totalrow], ignore_index=True)

# Look at it
print(t1.to_string(index=False))

# Hold onto this table; we will add it to the Word document later.


# Table 2: Descriptive Stats ---------------------------------------------------

# Select the variables to include in the table and apply variable labels.
# This is the Python equivalent of select(...) |> label_to_colnames().
descripdata = regdata[["roa_lead_1", "roa", "loss", "rd", "at", "mve"]].rename(
    columns=var_labels
)

# Compute descriptive statistics.
# pandas describe() gives count, mean, std, min, 25%, 50%, 75%, max.
# We rebuild this manually to match the modelsummary/datasummary output exactly.
def desc_stats(col):
    """Return a Series of formatted descriptive statistics for one column."""
    x = col.dropna()
    return pd.Series({
        "N":      f"{len(x):,.0f}",
        "Mean":   f"{x.mean():,.3f}",
        "SD":     f"{x.std():,.3f}",
        "Min":    f"{x.min():,.3f}",
        "P25":    f"{x.quantile(0.25):,.3f}",
        "Median": f"{x.median():,.3f}",
        "P75":    f"{x.quantile(0.75):,.3f}",
        "Max":    f"{x.max():,.3f}",
    })

t2 = (
    descripdata
    .apply(desc_stats)
    .T
    .reset_index()
    .rename(columns={"index": " "})
)

print(t2.to_string(index=False))


# Table 3: Correlation Matrix --------------------------------------------------

# Compute Pearson correlations for the upper triangle and Spearman for the
# lower triangle. This is the Python equivalent of R's
# datasummary_correlation(descripdata, method = "pearspear").

pearson = descripdata.corr(method="pearson")
spearman = descripdata.corr(method="spearman")

# Build combined matrix: Pearson above diagonal, Spearman below
corrmatrix = pearson.copy()
cols = corrmatrix.columns
n = len(cols)
for i in range(n):
    for j in range(n):
        if i == j:
            corrmatrix.iloc[i, j] = np.nan  # blank diagonal
        elif i > j:
            corrmatrix.iloc[i, j] = spearman.iloc[i, j]  # Spearman below

# Format as strings (3 decimal places) for display, blank the diagonal
def fmt_corr(x):
    if pd.isna(x):
        return ""
    return f"{x:.3f}"

t3 = corrmatrix.applymap(fmt_corr)

# Add row labels (variable names) as the first column
t3.insert(0, " ", t3.index)
t3 = t3.reset_index(drop=True)

print(t3.to_string(index=False))


# Table 4: Regression Table ----------------------------------------------------

# Run the five models using pyfixest (the Python equivalent of R's fixest::feols).
# pyfixest uses R-like formula syntax: "y ~ x | fe1 + fe2"
# fixef.rm="both" (remove singletons) is not directly available in pyfixest,
# but singletons are handled implicitly by the demeaning algorithm.

# Cluster standard errors by gvkey and calyear (two-way clustering).
# Equivalent to R's vcov = ~ gvkey + calyear.
vcov_spec = {"CRV1": "gvkey + calyear"}

models = {
    "Base":         pf.feols("roa_lead_1 ~ roa",
                             data=regdata, vcov=vcov_spec),
    "No FE":        pf.feols("roa_lead_1 ~ roa * loss",
                             data=regdata, vcov=vcov_spec),
    "Year FE":      pf.feols("roa_lead_1 ~ roa * loss | calyear",
                             data=regdata, vcov=vcov_spec),
    "Two-Way FE":   pf.feols("roa_lead_1 ~ roa * loss | calyear + gvkey",
                             data=regdata, vcov=vcov_spec),
    "With Controls":pf.feols("roa_lead_1 ~ roa * loss + at + rd + mve | calyear + gvkey",
                             data=regdata, vcov=vcov_spec),
}

# Coefficient map: controls the order and labels of coefficients in the table.
# Only coefficients listed here will appear; others are silently dropped.
# This is the Python equivalent of R's coef_map argument in modelsummary.
coef_map = {
    "roa":       "ROA_t",
    "loss":      "LOSS",
    "roa:loss":  "ROA_t x LOSS",
}

# Helper to add significance stars based on p-value.
# Equivalent to R's stars = c('*' = .1, '**' = .05, '***' = .01).
def add_stars(pval):
    if pval < 0.01:
        return "***"
    elif pval < 0.05:
        return "**"
    elif pval < 0.1:
        return "*"
    return ""

def extract_model_results(model, model_name):
    """Extract formatted coefficient estimates and t-stats from a pyfixest model."""
    params = model.coef()
    tvals  = model.tstat()
    pvals  = model.pvalue()
    rows   = []
    for coef_name, label in coef_map.items():
        if coef_name in params.index:
            est   = params[coef_name]
            tstat = tvals[coef_name]
            stars = add_stars(pvals[coef_name])
            rows.append({"term": label,            model_name: f"{est:.3f}{stars}"})
            rows.append({"term": "",               model_name: f"({tstat:.2f})"})
        else:
            rows.append({"term": label,            model_name: ""})
            rows.append({"term": "",               model_name: ""})
    return rows

# Build the coefficient section of the regression table
rows_list = []
for name, m in models.items():
    rows_list.append(extract_model_results(m, name))

# Transpose: we want one row per coefficient, one column per model
# Start from the first model's row structure, then fill in other models
first_model_name = list(models.keys())[0]
table_rows = extract_model_results(list(models.values())[0], first_model_name)

# Fill in the remaining models
for name, m in list(models.items())[1:]:
    model_rows = extract_model_results(m, name)
    for i, row in enumerate(model_rows):
        table_rows[i][name] = row[name]

t4_coefs = pd.DataFrame(table_rows)

# Goodness-of-fit rows
def has_fe(model, fe_name):
    """Check if a model includes a specific fixed effect."""
    try:
        return fe_name in model._fixef
    except Exception:
        return False

def has_controls(model):
    """Check if a model includes the control variables."""
    controls = ["at", "rd", "mve"]
    return all(c in model.coef().index for c in controls)

gof_data = []
for gof_label, getter in [
    ("Year FE",  lambda m: "X" if has_fe(m, "calyear") else ""),
    ("Firm FE",  lambda m: "X" if has_fe(m, "gvkey") else ""),
    ("N",        lambda m: f"{int(m.nobs):,.0f}"),
    ("R^2",      lambda m: f"{m.rsquared:.3f}"),
]:
    row = {"term": gof_label}
    for name, m in models.items():
        row[name] = getter(m)
    gof_data.append(row)

t4_gof = pd.DataFrame(gof_data)

# Combine coefficient rows and goodness-of-fit rows
col_order = ["term"] + list(models.keys())
t4 = pd.concat([t4_coefs, t4_gof])[col_order]

print(t4.to_string(index=False))


# Build a Word Document with all results ---------------------------------------

# These commands use the python-docx package, which is the Python equivalent
# of R's officer package. See https://python-docx.readthedocs.io/ for more.

def df_to_docx_table(doc, df, style="Table Grid"):
    """Add a pandas DataFrame as a formatted Word table to the document."""
    # Add table with a header row
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = style
    # Write header
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    # Write data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val) if not pd.isna(val) else ""
    return table

doc = Document()

# Table 1: Sample Frequency
doc.add_heading("Sample Frequency", level=1)
doc.add_paragraph("")
df_to_docx_table(doc, t1)
doc.add_page_break()

# Table 2: Descriptive Statistics
doc.add_heading("Descriptive Statistics", level=1)
doc.add_paragraph("")
df_to_docx_table(doc, t2)
doc.add_page_break()

# Table 3: Correlation Matrix
doc.add_heading("Correlation Matrix", level=1)
doc.add_paragraph("")
df_to_docx_table(doc, t3)
doc.add_page_break()

# Table 4: Regression Table
doc.add_heading("Regression Table", level=1)
doc.add_paragraph("")
df_to_docx_table(doc, t4)
doc.add_page_break()

# Figure 1: Losses by FF12 Industry
# Equivalent to R's body_add_img()
doc.add_heading("Figure 1", level=1)
doc.add_paragraph("")
doc.add_picture(f"{data_dir}/output/ff12_fig.png",
                width=Inches(4.2), height=Inches(3.6))

# Save the document.
# f-strings are the Python equivalent of R's glue().
doc.save(f"{data_dir}/output/tables-py.docx")

print(f"Word document saved to {data_dir}/output/tables-py.docx")

# Since these are plain Word tables, they can also be styled further using
# python-docx's table styling API. See https://python-docx.readthedocs.io/
# for the full range of options.
